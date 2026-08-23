#!/usr/bin/env python3
"""Assemble captured DocSend page images into Word notebooks.

Reads the manifest.json written by capture_dataroom.py and produces one .docx
per source document (a "notebook"), each page image on its own Word page, sized
to fill it. Orientation follows the pages themselves, so 16:9 decks come out
landscape instead of letterboxed on portrait paper.

Images are downscaled and re-encoded as JPEG before they go in, because raw
DocSend PNGs are often 3-5x larger than they need to be for reading and
annotating. --no-recompress keeps the originals when you need archival quality.

Usage:
    python3 build_docx.py <capture-dir> [--out <dir>] [options]
"""

from __future__ import annotations

import argparse
import json
import shutil
import sys
import tempfile
from pathlib import Path

try:
    from docx import Document
    from docx.enum.section import WD_ORIENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Emu, Inches, Pt, RGBColor
except ImportError:  # pragma: no cover - environment guard
    sys.exit("python-docx is not installed. Run: python3 -m pip install python-docx pillow")

try:
    from PIL import Image
except ImportError:  # Pillow is optional; without it we skip recompression.
    Image = None


def log(msg: str) -> None:
    print(msg, flush=True)


def human(n: float) -> str:
    for unit in ("B", "KB", "MB", "GB"):
        if n < 1024 or unit == "GB":
            return f"{n:.0f}{unit}" if unit == "B" else f"{n:.1f}{unit}"
        n /= 1024.0
    return f"{n:.1f}GB"


def image_size(path: Path) -> tuple[int, int] | None:
    """Pixel dimensions, via Pillow if present, else python-docx's own reader."""
    if Image is not None:
        try:
            with Image.open(path) as im:
                return im.size
        except Exception:
            return None
    try:
        from docx.image.image import Image as DocxImage

        img = DocxImage.from_file(str(path))
        return img.px_width, img.px_height
    except Exception:
        return None


def prepare(src: Path, work: Path, max_width: int, quality: int, recompress: bool) -> Path:
    """Return the file to embed: recompressed copy, or the original."""
    if not recompress or Image is None:
        return src
    try:
        with Image.open(src) as im:
            im = im.convert("RGB")
            if im.width > max_width:
                height = round(im.height * max_width / im.width)
                im = im.resize((max_width, height), Image.LANCZOS)
            out = work / (src.stem + ".jpg")
            im.save(out, "JPEG", quality=quality, optimize=True, progressive=True)
        return out if out.stat().st_size < src.stat().st_size else src
    except Exception:
        return src


def new_document(landscape: bool, margin_in: float) -> "Document":
    doc = Document()
    section = doc.sections[0]
    if landscape and section.page_width < section.page_height:
        section.orientation = WD_ORIENT.LANDSCAPE
        section.page_width, section.page_height = section.page_height, section.page_width
    for attr in ("top_margin", "bottom_margin", "left_margin", "right_margin"):
        setattr(section, attr, Inches(margin_in))
    return doc


def fitted_size(section, px: tuple[int, int] | None, caption_space: Emu) -> dict:
    """Scale to fill the text area without spilling onto the next Word page."""
    avail_w = section.page_width - section.left_margin - section.right_margin
    avail_h = section.page_height - section.top_margin - section.bottom_margin - caption_space
    if not px or px[0] <= 0 or px[1] <= 0:
        return {"width": avail_w}
    if px[1] / px[0] * avail_w > avail_h:
        return {"height": Emu(int(avail_h))}
    return {"width": Emu(int(avail_w))}


def add_caption(doc, text: str) -> None:
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.paragraph_format.space_before = Pt(2)
    para.paragraph_format.space_after = Pt(0)
    run = para.add_run(text)
    run.font.size = Pt(7)
    run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)


def build_one(doc_record: dict, capture_dir: Path, out_path: Path, args, manifest: dict) -> dict:
    pages = doc_record.get("pages") or []
    if not pages:
        return {"title": doc_record.get("title"), "pages": 0, "path": None, "bytes": 0}

    page_dir = capture_dir / doc_record["dir"]
    first = image_size(page_dir / pages[0]["file"])
    landscape = bool(first and first[0] > first[1])
    docx = new_document(landscape, args.margin)
    section = docx.sections[0]
    caption_space = Inches(0.18) if args.captions else Emu(0)

    docx.core_properties.title = doc_record.get("title") or out_path.stem
    docx.core_properties.comments = (
        f"Captured from {doc_record.get('url', '')} on {manifest.get('captured_at', '')}"
    )

    work = Path(tempfile.mkdtemp(prefix="docsend-docx-"))
    try:
        for i, page in enumerate(pages):
            src = page_dir / page["file"]
            if not src.exists():
                log(f"      missing {src.name}, skipping")
                continue
            embed = prepare(src, work, args.max_width_px, args.quality, not args.no_recompress)
            if i:
                docx.add_page_break()
            docx.add_picture(str(embed), **fitted_size(section, image_size(embed), caption_space))
            docx.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
            if args.captions:
                add_caption(docx, f"{doc_record.get('title', '')} — p.{page['n']}")
            if embed != src:
                embed.unlink(missing_ok=True)  # keep peak disk/memory flat
        out_path.parent.mkdir(parents=True, exist_ok=True)
        docx.save(str(out_path))
    finally:
        shutil.rmtree(work, ignore_errors=True)

    size = out_path.stat().st_size
    log(f"  {out_path.name}  ({len(pages)} pages, {human(size)})")
    return {"title": doc_record.get("title"), "pages": len(pages), "path": out_path, "bytes": size}


def build_index(results: list[dict], manifest: dict, out_path: Path) -> None:
    docx = Document()
    docx.add_heading(manifest.get("room_title") or "DocSend capture", level=1)
    meta = docx.add_paragraph()
    meta.add_run(
        f"Source: {manifest.get('source_url', 'n/a')}\n"
        f"Captured: {manifest.get('captured_at', 'n/a')} as {manifest.get('viewer_email', 'n/a')}\n"
        f"Documents: {len(results)} · Pages: {sum(r['pages'] for r in results)}"
    ).font.size = Pt(9)
    table = docx.add_table(rows=1, cols=3)
    table.style = "Light Grid Accent 1"
    for cell, text in zip(table.rows[0].cells, ("Document", "Pages", "File")):
        cell.text = text
    for r in results:
        row = table.add_row().cells
        row[0].text = str(r["title"] or "")
        row[1].text = str(r["pages"])
        row[2].text = r["path"].name if r["path"] else "(no pages captured)"
    docx.save(str(out_path))
    log(f"  {out_path.name}  (index)")


def main(argv=None) -> int:
    p = argparse.ArgumentParser(description=__doc__, formatter_class=argparse.RawDescriptionHelpFormatter)
    p.add_argument("capture_dir", help="Directory produced by capture_dataroom.py (holds manifest.json)")
    p.add_argument("--out", default=None, help="Where to write .docx files (default: <capture-dir>/word)")
    p.add_argument("--mode", choices=["per-doc", "single"], default="per-doc",
                   help="One notebook per document (default) or all documents in one file")
    p.add_argument("--max-width-px", type=int, default=1600, help="Downscale wider images (default: 1600)")
    p.add_argument("--quality", type=int, default=80, help="JPEG quality (default: 80)")
    p.add_argument("--no-recompress", action="store_true", help="Embed original image bytes untouched")
    p.add_argument("--margin", type=float, default=0.3, help="Page margin in inches (default: 0.3)")
    p.add_argument("--no-captions", dest="captions", action="store_false", help="Omit the page-number captions")
    p.add_argument("--no-index", dest="index", action="store_false", help="Skip the index document")
    args = p.parse_args(argv)

    capture_dir = Path(args.capture_dir).expanduser().resolve()
    manifest_path = capture_dir / "manifest.json"
    if not manifest_path.exists():
        sys.exit(f"No manifest.json in {capture_dir}. Run capture_dataroom.py first.")
    manifest = json.loads(manifest_path.read_text())
    docs = [d for d in manifest.get("documents", []) if d.get("pages")]
    if not docs:
        sys.exit("The manifest has no captured pages. Check the capture output for errors.")

    # A capture run with --discard-pages keeps only the collated PDFs, so say that
    # plainly instead of printing one "missing file" line per page.
    present = sum(
        1 for d in docs for pg in d["pages"] if (capture_dir / d["dir"] / pg["file"]).exists()
    )
    if not present:
        collated = [d for d in docs if d.get("pdf")]
        if collated:
            sys.exit(
                f"The page images are gone from {capture_dir} - this capture was run with\n"
                f"--discard-pages, so the {len(collated)} collated PDF(s) in {capture_dir / 'pdf'}\n"
                "are the copy. Word notebooks need the page images: re-capture without\n"
                "--discard-pages if you want them."
            )
        sys.exit(f"No page images found under {capture_dir}. Re-run the capture.")

    out_dir = Path(args.out).expanduser().resolve() if args.out else capture_dir / "word"
    out_dir.mkdir(parents=True, exist_ok=True)
    log(f"Building Word notebooks in {out_dir}")

    if args.mode == "single":
        room = (manifest.get("room_title") or "dataroom").strip()[:60] or "dataroom"
        safe = " ".join("".join(c for c in room if c.isalnum() or c in " -_.").split()) or "dataroom"
        # Date-stamp the combined file: one file per room per capture is far
        # easier to keep straight than a single name that silently gets replaced.
        stamp = (manifest.get("captured_at") or "")[:10] or "undated"
        results = [build_single(docs, capture_dir, out_dir / f"{safe} {stamp}.docx", args, manifest)]
    else:
        results = []
        for d in docs:
            name = f"{d['dir']}.docx"
            results.append(build_one(d, capture_dir, out_dir / name, args, manifest))
        if args.index:
            build_index(results, manifest, out_dir / "00-INDEX.docx")

    total = sum(r["bytes"] for r in results if r["path"])
    log("")
    log(f"{len([r for r in results if r['path']])} file(s), {sum(r['pages'] for r in results)} pages, {human(total)} total")
    log(f"Open them from: {out_dir}")
    return 0


def build_single(docs: list[dict], capture_dir: Path, out_path: Path, args, manifest: dict) -> dict:
    """All documents in one .docx, each starting on a fresh page with a heading."""
    first = image_size(capture_dir / docs[0]["dir"] / docs[0]["pages"][0]["file"])
    landscape = bool(first and first[0] > first[1])
    docx = new_document(landscape, args.margin)
    section = docx.sections[0]
    caption_space = Inches(0.18) if args.captions else Emu(0)
    docx.core_properties.title = manifest.get("room_title") or out_path.stem

    work = Path(tempfile.mkdtemp(prefix="docsend-docx-"))
    written = 0
    started = False
    try:
        for d in docs:
            # A divider page per document: cheap, and it makes a merged file
            # navigable in Word's outline instead of one undifferentiated run.
            if started:
                docx.add_page_break()
            heading = docx.add_heading(d.get("title") or d["dir"], level=1)
            heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
            started = True
            for page in d["pages"]:
                src = capture_dir / d["dir"] / page["file"]
                if not src.exists():
                    continue
                docx.add_page_break()
                embed = prepare(src, work, args.max_width_px, args.quality, not args.no_recompress)
                docx.add_picture(str(embed), **fitted_size(section, image_size(embed), caption_space))
                docx.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
                if args.captions:
                    add_caption(docx, f"{d.get('title', '')} — p.{page['n']}")
                written += 1
                if embed != src:
                    embed.unlink(missing_ok=True)
        out_path.parent.mkdir(parents=True, exist_ok=True)
        docx.save(str(out_path))
    finally:
        shutil.rmtree(work, ignore_errors=True)
    size = out_path.stat().st_size
    log(f"  {out_path.name}  ({written} pages, {human(size)})")
    return {"title": docx.core_properties.title, "pages": written, "path": out_path, "bytes": size}


if __name__ == "__main__":
    sys.exit(main())
